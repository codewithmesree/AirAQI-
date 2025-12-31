import os
import csv
from datetime import datetime
from fpdf import FPDF
from docx import Document

REPORT_DIR = os.path.join(os.path.dirname(__file__), 'static', 'reports')
os.makedirs(REPORT_DIR, exist_ok=True)

class PDFReport(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(80)
        self.cell(30, 10, 'AirAQI Quality Report', 0, 0, 'C')
        self.ln(20)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, 'Page ' + str(self.page_no()) + '/{nb}', 0, 0, 'C')

def generate_report(report_name, file_format, firebase_uid):
    """
    Generates a report file (PDF, CSV, or DOCX) with mock analysis data.
    Returns the filename relative to static/reports.
    """
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"report_{firebase_uid[:5]}_{timestamp}.{file_format.lower()}"
    filepath = os.path.join(REPORT_DIR, filename)

    # Mock Data for Report
    data = [
        ['Date', 'AQI', 'Status', 'Main Pollutant'],
        ['2023-10-01', '150', 'Unhealthy', 'PM2.5'],
        ['2023-10-02', '180', 'Unhealthy', 'PM2.5'],
        ['2023-10-03', '210', 'Very Unhealthy', 'PM10'],
        ['2023-10-04', '90', 'Moderate', 'NO2'],
        ['2023-10-05', '45', 'Good', 'O3'],
    ]

    try:
        if file_format.lower() == 'pdf':
            pdf = PDFReport()
            pdf.alias_nb_pages()
            pdf.add_page()
            pdf.set_font('Arial', '', 12)
            
            pdf.cell(0, 10, f'Report Name: {report_name}', 0, 1)
            pdf.cell(0, 10, f'Generated On: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', 0, 1)
            pdf.ln(10)
            
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 10, 'Weekly Analysis Data:', 0, 1)
            pdf.set_font('Arial', '', 12)
            
            # Simple Table
            col_width = 40
            th = pdf.font_size
            
            for row in data:
                for datum in row:
                    pdf.cell(col_width, 2 * th, str(datum), border=1)
                pdf.ln(2 * th)
                
            pdf.ln(10)
            pdf.multi_cell(0, 10, 'Summary: Air quality has shown significant fluctuation over the past week. Peak pollution was observed on Oct 3rd due to stable atmospheric conditions.')
            
            pdf.output(filepath)

        elif file_format.lower() == 'csv':
            with open(filepath, 'w', newline='') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerows(data)

        elif file_format.lower() == 'docx': # Word
            document = Document()
            document.add_heading('AirAQI Quality Report', 0)
            
            document.add_paragraph(f'Report Name: {report_name}')
            document.add_paragraph(f'Generated On: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            
            document.add_heading('Weekly Analysis Data', level=1)
            
            table = document.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Date'
            hdr_cells[1].text = 'AQI'
            hdr_cells[2].text = 'Status'
            hdr_cells[3].text = 'Main Pollutant'
            
            for date, aqi, status, pollutant in data[1:]:
                row_cells = table.add_row().cells
                row_cells[0].text = date
                row_cells[1].text = aqi
                row_cells[2].text = status
                row_cells[3].text = pollutant
                
            document.add_paragraph('\nSummary: Air quality has shown significant fluctuation over the past week.')
            
            document.save(filepath)
            
        return filename

    except Exception as e:
        error_msg = f"Error generating report: {e}"
        print(error_msg)
        try:
             with open(os.path.join(os.path.dirname(__file__), 'gen_error.txt'), 'w') as f:
                 f.write(str(e))
                 import traceback
                 traceback.print_exc(file=f)
        except:
             pass
        return None
